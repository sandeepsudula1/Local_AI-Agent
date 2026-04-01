"""
Generate a professional Word document for the Local AI Assistant project.
Run: python generate_project_doc.py
Output: Local_AI_Assistant_Project_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ───────────────────────────────────────────────────────────
C_NAVY    = RGBColor(0x1F, 0x35, 0x64)   # dark navy   — headings
C_BLUE    = RGBColor(0x1E, 0x6B, 0xB8)   # mid blue    — sub-headings
C_TEAL    = RGBColor(0x00, 0x7A, 0x87)   # teal        — accent
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_LGRAY   = RGBColor(0xF2, 0xF2, 0xF2)   # light-grey  — table header bg
C_DGRAY   = RGBColor(0x40, 0x40, 0x40)   # body text
C_GREEN   = RGBColor(0x10, 0x7C, 0x10)
C_ORANGE  = RGBColor(0xCC, 0x55, 0x00)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Thin borders on every cell."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "CCCCCC")
                tcBorders.append(el)
            tcPr.append(tcBorders)

def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = color or C_NAVY
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F3564")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = color or C_BLUE
    else:
        run.font.size  = Pt(12)
        run.font.color.rgb = color or C_TEAL
    return p

def add_body(text, indent=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.font.color.rgb = color or C_DGRAY
    return p

def add_bullet(text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or C_DGRAY
    return p

def add_two_col_table(rows, header=None):
    cols = 2
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    idx = 0
    if header:
        for j, h in enumerate(header):
            cell = table.cell(0, j)
            set_cell_bg(cell, "1E6BB8")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = C_WHITE
            run.font.size = Pt(10)
        idx = 1
    for r, (a, b) in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else "F5F8FF"
        for j, val in enumerate([a, b]):
            cell = table.cell(r + idx, j)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = C_DGRAY
    set_cell_borders(table)
    doc.add_paragraph()

def add_wide_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_bg(cell, "1F3564")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(9.5)
    for r, row_data in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else "F2F6FC"
        for j, val in enumerate(row_data):
            cell = table.cell(r + 1, j)
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = C_DGRAY
    set_cell_borders(table)
    doc.add_paragraph()

def add_box(title, content_lines, box_color="EAF4FB", border_color="1E6BB8"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, box_color)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "12")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1E, 0x6B, 0xB8)
    for line in content_lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(1)
        run2 = p2.add_run(line)
        run2.font.size = Pt(10)
        run2.font.color.rgb = C_DGRAY
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("LOCAL AI ASSISTANT")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = C_NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Multi-Agent Offline AI System")
r2.font.size = Pt(18)
r2.font.color.rgb = C_BLUE

doc.add_paragraph()
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
line.paragraph_format.space_after = Pt(2)
lr = line.add_run("─" * 60)
lr.font.color.rgb = C_TEAL

for label, value in [
    ("Prepared by", "Sandeep"),
    ("Date",        datetime.date.today().strftime("%d %B %Y")),
    ("Version",     "1.1"),
    ("Status",      "Production-Ready"),
    ("Classification", "Internal / Manager Review"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = C_NAVY
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_DGRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ════════════════════════════════════════════════════════════════════════════
add_heading("Table of Contents", level=1)
toc_entries = [
    ("1.", "Executive Summary"),
    ("2.", "Project Overview"),
    ("3.", "System Architecture"),
    ("4.", "Technologies & Stack"),
    ("5.", "Core AI Capabilities & Automation"),
    ("6.", "Module-by-Module Implementation"),
    ("7.", "Data Flow & Process"),
    ("8.", "Security & Access Control"),
    ("9.", "User Interfaces"),
    ("10.", "Testing & Quality"),
    ("11.", "Benefits & Business Value"),
    ("12.", "Architecture Mind Map"),
    ("13.", "Future Roadmap"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"  {num}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = C_BLUE
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = C_DGRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_heading("1. Executive Summary", level=1)
add_body(
    "The Local AI Assistant (v1.1) is a fully offline, privacy-first, multi-agent "
    "artificial intelligence platform built to automate everyday knowledge-work tasks "
    "directly on a local Windows machine — without sending any data to the cloud. "
    "The system understands natural-language requests from users and intelligently "
    "routes them to specialised AI agents that handle documents, emails, reminders, "
    "audio, and general conversation."
)
add_body(
    "At its core, the platform combines a locally running Large Language Model (LLM) "
    "via Ollama, a semantic vector database (ChromaDB), sentence-level embeddings "
    "(Sentence-Transformers), and a layered multi-agent orchestration pipeline. "
    "It is accessible through a command-line interface (CLI) and a modern browser-based "
    "chat UI built with Streamlit, and exposes all capabilities via a Model Context "
    "Protocol (MCP) server for integration with external AI clients such as Claude "
    "Desktop and VS Code Copilot extensions."
)

add_box("Key Highlights", [
    "✔  100% offline — no cloud API, no data egress",
    "✔  Natural-language interface for documents, email, reminders, and audio",
    "✔  Semantic search over local files using RAG (Retrieval-Augmented Generation)",
    "✔  AI-generated email replies with draft review and safe two-step sending",
    "✔  Audio transcription and semantic Q&A with timestamps",
    "✔  Persistent conversation memory with session context",
    "✔  Role-based folder access control with runtime permission management",
    "✔  MCP server — plug into Claude Desktop / VS Code Copilot",
    "✔  Streamlit web UI + CLI",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  2. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
add_heading("2. Project Overview", level=1)
add_heading("2.1  Problem Statement", level=2)
add_body(
    "Office workers and developers routinely lose time switching between tools to search "
    "documents, manage emails, set reminders, and recall past conversations. "
    "Cloud-based AI assistants (e.g., ChatGPT, Copilot) raise data-privacy concerns "
    "and require internet connectivity. There was no single, privacy-safe, locally "
    "operating assistant capable of performing all these tasks from one interface."
)

add_heading("2.2  Solution", level=2)
add_body(
    "The Local AI Assistant solves this by providing a unified natural-language interface "
    "that runs entirely on the local machine. Users converse naturally; the system "
    "classifies intent, selects the right specialised agent, and returns a precise, "
    "context-aware answer — all without any internet connection or data leaving the device."
)

add_heading("2.3  Scope of Functionality", level=2)
add_wide_table(
    ["Domain", "What the system can do", "AI Mechanism"],
    [
        ["Documents",  "Search, summarise, list, compare files from any authorised folder",                     "RAG + ChromaDB + LLM"],
        ["Email",      "Search inbox, summarise, generate AI replies, draft review, send via SMTP",             "Semantic search + LLM generation"],
        ["Reminders",  "Set, list, delete scheduled reminders; Windows toast notifications",                    "NLP date parsing + background thread"],
        ["Audio",      "Transcribe audio files, index transcript, answer questions with timestamps",             "Whisper ASR + ChromaDB + LLM"],
        ["General",    "Answer open questions, compare topics, programming help, general knowledge",             "Local LLM (Ollama)"],
        ["Memory",     "Remember user name, preferences, last files; context follow-up across turns",           "In-memory + JSON persistence"],
        ["Access Ctrl","Grant/deny folder access at runtime; normalize paths; prevent unauthorised retrieval",   "Rule-based + permission store"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  3. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
add_heading("3. System Architecture", level=1)
add_heading("3.1  Layered Architecture Overview", level=2)
add_body(
    "The system is divided into six logical layers, each with a clear responsibility. "
    "Every user request flows top-down through these layers and the response travels back up."
)

layers = [
    ("Layer 1 — User Interface",
     "CLI (main.py) and Streamlit Web UI (streamlit_app.py). Accepts natural-language input; renders structured responses with intent labels, tool names, latency, and source attribution."),
    ("Layer 2 — Orchestrator (pipelines/orchestrator.py)",
     "Central coordinator. Manages the full request lifecycle: access control → intent classification → tool routing → tool execution → memory update → response formatting."),
    ("Layer 3 — Intent Classification (core/intent_classifier.py)",
     "Hybrid classifier: deterministic regex fast-path → email context guardrails → LLM classification (Ollama, format=JSON) → heuristic fallback → planner_agent fallback."),
    ("Layer 4 — Routing & Tool Execution (core/router.py, core/tool_executor.py)",
     "Router maps intent labels to canonical tool names. ToolExecutor invokes the correct agent or service and returns a ToolResult."),
    ("Layer 5 — Agents (agents/)",
     "Specialised agents: retrieval, summary, topic, email-query, email-reply, email-summarizer, reminder, audio, general, planner."),
    ("Layer 6 — Engines & Services",
     "RAG engine (ChromaDB + embeddings), vector store service, document service, reminder service, email send service, draft manager, document indexer service, audio agent."),
]
for title, desc in layers:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.4)
    r1 = p.add_run(f"  {title}:  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = C_BLUE
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = C_DGRAY

add_heading("3.2  Key Architectural Decisions", level=2)
decisions = [
    ("Offline-first by design",       "All LLM calls go to Ollama (localhost). No outbound HTTP to AI APIs."),
    ("Lazy service initialisation",    "Vector DB, email cache, and reminder service start in background threads so the CLI is instantly responsive."),
    ("MCP protocol support",           "FastMCP server exposes all tools over stdio or SSE for integration with any MCP-compatible client."),
    ("Deterministic intent pipeline",  "Regex guardrails prevent LLM misclassification for high-frequency patterns (email reply, document listing)."),
    ("RAG folder scoping",             "All document retrieval is scoped to an authorised folder so the LLM never mixes data from different sources."),
    ("Draft-before-send email safety", "AI-generated replies are stored as drafts. The user must explicitly confirm sending — auto-send is architecturally blocked."),
    ("Persistent conversation memory", "Session facts (name, preferences, last file, last folder, last intent) survive the session for context-aware follow-ups."),
]
add_two_col_table(decisions, header=["Design Decision", "Rationale"])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  4. TECHNOLOGIES & STACK
# ════════════════════════════════════════════════════════════════════════════
add_heading("4. Technologies & Stack", level=1)

add_heading("4.1  Technology Matrix", level=2)
add_wide_table(
    ["Category", "Technology / Library", "Version", "Purpose"],
    [
        ["LLM Runtime",       "Ollama",                              "0.6.1",    "Local LLM server — runs llama3.2:1b on CPU/GPU"],
        ["LLM Model",         "Llama 3.2 1B (Meta)",                "3.2",      "Natural language understanding, generation, classification"],
        ["Speech-to-Text",    "OpenAI Whisper (via faster-whisper)", "base",     "Audio transcription with timestamps"],
        ["Embeddings",        "Sentence-Transformers all-MiniLM-L6-v2","5.2.3", "384-dim dense vectors for semantic search"],
        ["Vector Database",   "ChromaDB",                           "1.5.0",    "Persistent vector store for documents, emails, audio"],
        ["LLM Framework",     "LangChain + LangChain-Community",    "0.1.12",   "Document loading, chunking, RAG pipeline"],
        ["Web UI",            "Streamlit",                          "≥1.35",    "Browser-based chat interface"],
        ["MCP Protocol",      "FastMCP (mcp SDK)",                  "latest",   "Tool exposure to Claude Desktop / VS Code Copilot"],
        ["Email — IMAP",      "IMAPClient",                         "3.1.0",    "Fetch emails from Gmail/Outlook via IMAP"],
        ["Email — Gmail API", "google-api-python-client",           "≥2.128",   "OAuth2 Gmail access"],
        ["Email — SMTP",      "smtplib (stdlib)",                   "built-in", "Send emails; TLS/SSL"],
        ["OCR",               "Tesseract + pytesseract",            "0.3.13",   "Extract text from images (PNG, JPG)"],
        ["Document Parsing",  "python-docx, python-pptx, openpyxl", "latest",   "Read Word, PowerPoint, Excel files"],
        ["PDF Parsing",       "PyPDFLoader (LangChain)",            "built-in", "Load and chunk PDF files"],
        ["Notifications",     "plyer + pywin32",                    "2.1.0",    "Windows toast notifications for reminders"],
        ["Date Parsing",      "dateparser",                          "1.3.0",    "Parse natural-language time expressions"],
        ["Configuration",     "pydantic-settings",                  "2.13.1",   "Typed, validated application settings"],
        ["Data Processing",   "pandas, numpy, scipy",               "latest",   "CSV parsing, numerical operations"],
        ["Testing",           "pytest + pytest-mock",               "≥7.0",     "Unit and integration tests"],
        ["Language",          "Python",                             "3.11",     "Core implementation language"],
        ["OS Target",         "Windows 10/11",                      "—",        "Primary deployment platform"],
    ]
)

add_heading("4.2  AI Models Used", level=2)
add_two_col_table([
    ("llama3.2:1b (Meta via Ollama)",      "Primary LLM — intent classification, answer generation, email reply drafting, general conversation. Runs fully on-device."),
    ("all-MiniLM-L6-v2 (Sentence-Transformers)", "Embedding model — converts text chunks and queries to 384-dimensional vectors for semantic similarity search."),
    ("OpenAI Whisper base",                "Automatic Speech Recognition — transcribes audio files (.mp3, .wav, .m4a) to text with word-level timestamps."),
], header=["Model", "Role"])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  5. CORE AI CAPABILITIES & AUTOMATION
# ════════════════════════════════════════════════════════════════════════════
add_heading("5. Core AI Capabilities & Automation", level=1)
add_body(
    "This section highlights the AI-driven automation features that differentiate "
    "this system from a simple script or chatbot."
)

add_heading("5.1  Retrieval-Augmented Generation (RAG)", level=2)
add_body(
    "RAG is the core document intelligence mechanism. When a user asks a question "
    "about their files, the system does NOT send the entire document to the LLM. Instead:"
)
steps_rag = [
    "The user's question is converted to a 384-dim embedding vector.",
    "The vector is compared against all pre-indexed document chunk vectors in ChromaDB via cosine similarity.",
    "The top-k most relevant chunks (threshold configurable) are retrieved.",
    "Only those chunks are injected into the LLM prompt as context.",
    "The LLM synthesises a grounded answer citing the source document.",
]
for i, s in enumerate(steps_rag, 1):
    add_bullet(f"Step {i}: {s}")

add_body(
    "This approach means the system can answer questions about multi-hundred-page "
    "documents in under 2 seconds while staying within the LLM's context window. "
    "Supported file formats: PDF, DOCX, PPTX, XLSX, CSV, TXT, MD, PNG, JPG (OCR)."
)

add_heading("5.2  Hybrid Intent Classification", level=2)
add_body(
    "Every user message passes through a 6-stage intent classification pipeline before "
    "any LLM call is made, ensuring fast, accurate, and consistent routing:"
)
pipeline_stages = [
    ("Stage 1 — Regex fast-path",     "Instant deterministic patterns for greetings, time/date, audio, document listing, email reply/send. No LLM needed."),
    ("Stage 2 — Email context override","When a prior email exists in session memory and a reply signal is detected, EMAIL_REPLY fires immediately regardless of other signals."),
    ("Stage 3 — Context guardrails",   "State-machine transitions: EMAIL_SEARCH → EMAIL_REPLY → EMAIL_SEND, driven by conversation history."),
    ("Stage 4 — LLM classification",   "Ollama called with format=json to enforce structured output. Free-text LLM responses are rejected (not keyword-scanned)."),
    ("Stage 5 — Heuristic fallback",   "Conservative keyword-based safety net fires only when LLM is unavailable or returns invalid JSON."),
    ("Stage 6 — Planner agent",        "Legacy regex-based planner_agent as final fallback before defaulting to GENERAL."),
]
add_two_col_table(pipeline_stages, header=["Stage", "Description"])

add_heading("5.3  AI-Generated Email Replies", level=2)
add_body(
    "The email reply subsystem automates the most time-consuming part of email management. "
    "The AI:"
)
email_steps = [
    "Locates the target email from inbox via semantic search (sender, subject, or context).",
    "Builds a grounded LLM prompt with strict anti-hallucination constraints (8 explicit negations).",
    "Generates a context-aware reply in the requested tone (professional, friendly, casual, formal).",
    "Saves it as a versioned draft (e.g., draft_20260331_001) with full audit metadata.",
    "Presents the draft to the user for review — no email is ever auto-sent.",
    "On explicit user confirmation ('send it' / 'yes'), sends via SMTP and marks draft as 'sent'.",
]
for s in email_steps:
    add_bullet(s)

add_heading("5.4  Audio Intelligence", level=2)
add_body(
    "Users can drop any audio file (MP3, WAV, M4A, FLAC) and the system will:"
)
audio_steps = [
    "Transcribe the audio using Whisper ASR, producing word-level timestamps.",
    "Chunk the transcript and index it in a dedicated ChromaDB collection.",
    "Allow semantic Q&A: 'What was discussed at 3:25?', 'Summarise the meeting'.",
    "Return answers with exact timestamp citations from the transcript.",
]
for s in audio_steps:
    add_bullet(s)

add_heading("5.5  Persistent Conversation Memory", level=2)
add_body(
    "The memory subsystem maintains context across the entire session, enabling "
    "natural multi-turn conversations:"
)
memory_items = [
    ("last_file",   "Most recently referenced document — enables 'summarise this file' without re-specifying."),
    ("last_folder", "Most recently used folder — enables 'show files from that folder' follow-ups."),
    ("last_intent", "Previous classified intent — feeds into the context guardrail state machine."),
    ("last_email",  "Last email shown in a search — enables 'reply to it' without re-searching."),
    ("User facts",  "Name, role, preferences extracted automatically from conversation and persisted to memory.json."),
]
add_two_col_table(memory_items, header=["Memory Slot", "Purpose"])

add_heading("5.6  Semantic Email Search", level=2)
add_body(
    "Email search goes beyond keyword matching. The system:"
)
for s in [
    "Converts the user's query to an embedding vector.",
    "Searches a ChromaDB collection of email subjects + bodies by cosine similarity.",
    "Combines semantic results (70% weight) with keyword fuzzy-match results (30% weight) using a hybrid scoring algorithm.",
    "Refreshes the live email cache from IMAP before every search (10-second TTL prevents redundant connections).",
    "Supports multi-field queries: 'urgent emails from Alice about the budget report'.",
]:
    add_bullet(s)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  6. MODULE-BY-MODULE IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════
add_heading("6. Module-by-Module Implementation", level=1)

modules = [
    ("core/intent_classifier.py",
     "Hybrid 6-stage intent classifier. Deterministic regex fast-path, email context state machine, LLM JSON classification, heuristic fallback. 700+ lines."),
    ("core/router.py",
     "Maps 30+ intent labels to canonical tool names (documents.search, email.reply, reminders.set, etc.). Merges static map with ToolCatalog."),
    ("core/tool_executor.py",
     "Runs the selected tool. Contains handlers for email reply/send (draft creation, SMTP), retrieval, summary, reminders, audio. 800+ lines."),
    ("core/access_control.py",
     "Access-control layer. Extracts Windows paths, normalises typos, checks permission store, issues ALLOW_FOLDER / REQUEST_PERMISSION / BLOCK / CLARIFY decisions."),
    ("core/permission_store.py",
     "Thread-safe store for granted/denied folder paths. Persists to data/granted_folders.json. Handles 5-minute permission expiry."),
    ("pipelines/orchestrator.py",
     "Central request orchestrator. 900+ lines. Manages the full pipeline from access control to memory update."),
    ("memory/conversation_memory.py",
     "Stores conversation history, user facts, last_file, last_folder, last_email, pending queries. Auto-extracts facts using regex."),
    ("agents/knowledge/retrieval_agent.py",
     "RAG-based document Q&A. Folder-scoped filtering, direct disk-first reading for named files, multi-intent dispatch."),
    ("agents/knowledge/email_query_agent.py",
     "Three-layer email loader (emails.json + email_cache.json + IMAP live). Hybrid semantic+keyword search with TTL caching."),
    ("agents/knowledge/email_reply_agent_v2.py",
     "LLM-based email reply generator. Anti-hallucination prompting, tone selection, 5-level email selection strategy."),
    ("agents/knowledge/audio_agent.py",
     "Whisper transcription, ChromaDB indexing, timestamped semantic Q&A over audio transcripts."),
    ("services/draft_manager.py",
     "Full email draft lifecycle. Thread-safe, JSON-persisted, versioned draft IDs, audit timestamps, status state machine."),
    ("services/vector_store_service.py",
     "Background-threaded ChromaDB loader. Auto-rebuilds when documents change. Staleness detection via file mtime."),
    ("services/document_indexer_service.py",
     "Dynamic folder indexer — indexes newly-granted folders on demand after permission is granted."),
    ("services/email_send_service.py",
     "SMTP email sender. Auto-detects provider (Gmail, Outlook, custom). TLS/SSL support. Email validation."),
    ("engines/rag_engine.py",
     "Core RAG pipeline. Embedding, retrieval, reranking, prompt building, LLM call. Used by retrieval and summary agents."),
    ("engines/embedding_engine.py",
     "Shared embedding utility. Singleton, lazy-load, batch processing, L2 normalisation."),
    ("agent_mcp/server.py",
     "FastMCP server. Registers 15 tools over stdio or SSE. Enables Claude Desktop and VS Code Copilot integration."),
    ("agent_mcp/bridge.py",
     "In-process MCP bridge. Routes intents to MCP tool wrappers without starting a server."),
    ("streamlit_app.py",
     "Streamlit Web UI. Chat bubbles, intent/tool/latency metadata per message, system status sidebar, memory panel."),
    ("main.py",
     "Production CLI entry point. Layered boot: logging → documents → vector store → reminders → email polling → CLI loop."),
]
add_wide_table(["Module", "Description"], modules)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  7. DATA FLOW & PROCESS
# ════════════════════════════════════════════════════════════════════════════
add_heading("7. Data Flow & Process", level=1)
add_heading("7.1  Complete Request Lifecycle", level=2)
add_body("Every user message follows this deterministic path:")

flow_steps = [
    ("1", "User Input",              "Text entered in CLI or Streamlit chat"),
    ("2", "Permission Intercept",    "Checks if input is a yes/no response to a pending permission request"),
    ("3", "Folder Clarification",    "Resolves pending folder-clarification turn if present"),
    ("4", "Access Control",          "Extracts path, classifies intent (content query vs. access check), returns ALLOW_FOLDER / REQUEST_PERMISSION / BLOCK / CLARIFY / PASS"),
    ("5", "Fact Extraction",         "Memory auto-extracts user facts from input"),
    ("6", "Intent Classification",   "6-stage hybrid pipeline → intent label (e.g., DOCUMENT_LIST, EMAIL_REPLY)"),
    ("7", "Intent Override",         "Orchestrator applies safety overrides (folder active → document intent; doc keywords + no email → not EMAIL_SEARCH)"),
    ("8", "Tool Routing",            "Router maps intent → tool name (documents.list, email.reply, reminders.set, etc.)"),
    ("9", "Tool Execution",          "ToolExecutor calls the agent / service with folder_path in context; logs selected folder"),
    ("10","Response Building",       "Tool result formatted → bullets generated → AgentResponse constructed"),
    ("11","Memory Update",           "last_intent, last_file, last_folder, last_email updated in session memory"),
    ("12","Structured Log",          "AgentLogger writes JSON log line with query, intent, tool, latency, source, error"),
    ("13","Output",                  "Answer displayed with intent label, tool, latency, source in UI"),
]
add_wide_table(["Step", "Stage", "Description"], flow_steps)

add_heading("7.2  Document Query Flow (RAG)", level=2)
rag_flow = [
    "User: 'What are the key findings in the Q4 report?'",
    "Access control: path extracted / permission checked → ALLOW_FOLDER (e.g., C:\\AI_Test_Documents)",
    "Intent: RETRIEVAL (or DOCUMENT_FOLDER_QUERY)",
    "Router: documents.search",
    "RAG Engine: query embedded → cosine search ChromaDB → top-3 chunks retrieved",
    "LLM: answer generated using ONLY retrieved chunks",
    "Response: answer + source filename displayed",
    "Memory: last_file = 'Q4_Report.pdf', last_folder = 'C:\\AI_Test_Documents'",
]
for s in rag_flow:
    add_bullet(s)

add_heading("7.3  Email Reply & Send Flow", level=2)
email_flow = [
    "User: 'Find emails from Alice about the project'  →  EMAIL_SEARCH → results shown + stored in memory",
    "User: 'Reply to it in professional tone'  →  EMAIL_REPLY (context override fires) → draft created",
    "Draft displayed: To, Subject, Body, Draft ID shown — email NOT sent",
    "User: 'Send it'  →  EMAIL_SEND → draft retrieved from draft_manager",
    "SMTP connection opened → email delivered → draft status = 'sent'",
    "Confirmation message shown with draft ID",
]
for s in email_flow:
    add_bullet(s)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  8. SECURITY & ACCESS CONTROL
# ════════════════════════════════════════════════════════════════════════════
add_heading("8. Security & Access Control", level=1)
add_body(
    "The access-control layer was designed to prevent unauthorised file access "
    "while remaining frictionless for the user."
)
sec_features = [
    ("Folder whitelist",          "Only folders explicitly granted by the user or pre-seeded in ALLOWED_FOLDERS can be queried. All other paths are blocked."),
    ("Runtime permission grants", "First-time access to a new folder triggers a permission prompt. The user types 'yes' to grant or 'no' to deny. Grants persist to data/granted_folders.json."),
    ("5-minute expiry",           "Unanswered permission requests expire after 5 minutes (permission_store) to prevent stale permission windows."),
    ("System path blocking",      "Root drives (C:\\), Windows, System32, Program Files, and AppData are hard-blocked and can never be requested via permission."),
    ("Path normalisation",        "Typos fixed: C:AI_Test → C:\\AI_Test; forward slashes converted; OneDrive Desktop auto-corrected."),
    ("Folder scoping",            "All RAG and document-list operations are scoped to the granted folder_path — the LLM never sees out-of-scope files."),
    ("Email passthrough",         "Email reply/response phrases bypass access control entirely (dedicated regex guard) to prevent false BLOCKs."),
    ("No cloud egress",           "All LLM calls go to localhost:11434 (Ollama). No query or document data is sent externally."),
]
add_two_col_table(sec_features, header=["Feature", "Detail"])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  9. USER INTERFACES
# ════════════════════════════════════════════════════════════════════════════
add_heading("9. User Interfaces", level=1)

add_heading("9.1  Command-Line Interface (CLI)", level=2)
add_body("Launch: python main.py  |  venv311\\Scripts\\activate → python main.py")
cli_features = [
    "Interactive read-eval-print loop (REPL)",
    "Built-in commands: 'list tools', 'forget everything', 'what do you remember', 'show permissions'",
    "Reminder confirmation workflow",
    "Email permission grant/deny workflow interception",
    "Auto-clears __pycache__ on startup so code changes take effect immediately",
]
for f in cli_features:
    add_bullet(f)

add_heading("9.2  Streamlit Web UI", level=2)
add_body("Launch: streamlit run streamlit_app.py")
ui_features = [
    "Modern chat bubble interface with full conversation history",
    "Per-message metadata ribbon: intent label | tool used | latency (ms) | source document",
    "Sidebar: system status panel (Ollama, Vector Store, Documents, Reminders)",
    "Sidebar: memory panel showing recalled facts",
    "Clear Memory button, tool catalogue viewer",
    "Services cached via @st.cache_resource — start exactly once per process",
    "Background email polling daemon (60-second intervals)",
]
for f in ui_features:
    add_bullet(f)

add_heading("9.3  MCP Server", level=2)
add_body("Launch: python main.py --mcp  (stdio) or  python main.py --mcp-sse  (HTTP/SSE on port 8765)")
add_body("Exposes 15 tools to any MCP-compatible AI client:")
mcp_tools = [
    ("reminders.set / .list / .delete",      "Create, list, and delete reminders from natural language"),
    ("email.search / .summarize / .list_all","Search inbox, summarise emails, list raw emails"),
    ("documents.search / .summarize",        "RAG search + document summarisation"),
    ("documents.topics / .list",             "Topic extraction and file listing"),
    ("audio.transcribe / .query / .list",    "Audio transcription, Q&A, file listing"),
    ("system.chat / .intent / .status",      "LLM conversation, intent inspection, health check"),
]
add_two_col_table(mcp_tools, header=["MCP Tool", "Capability"])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  10. TESTING & QUALITY
# ════════════════════════════════════════════════════════════════════════════
add_heading("10. Testing & Quality", level=1)
add_heading("10.1  Test Suite Results", level=2)
add_wide_table(
    ["Test Suite", "File", "Tests", "Status"],
    [
        ["Draft Flow Tests",              "scripts/test_draft_flow.py",                    "7/7",   "✅ PASS"],
        ["Conversation Memory Tests",     "test_email_context_debug.py",                   "4/4",   "✅ PASS"],
        ["Intent Classifier Tests",       "test_intent_classifier.py",                     "5/5",   "✅ PASS"],
        ["Orchestrator History Tests",    "test_email_context_debug.py",                   "4/4",   "✅ PASS"],
        ["Email Workflow Integration",    "scripts/test_email_draft_send_integration.py",  "4/4",   "✅ PASS"],
        ["Email Reply Tests",             "scripts/test_email_reply.py",                   "8 tests","✅ PASS"],
        ["Simple Intent Tests",           "test_intent_simple.py",                         "—",     "✅ PASS"],
        ["Import Sanity Tests",           "test_imports.py",                               "—",     "✅ PASS"],
        ["Draft Persistence Tests",       "test_draft_persistence.py",                     "—",     "✅ PASS"],
    ]
)

add_heading("10.2  Quality Practices", level=2)
quality = [
    "Type annotations (Python typing module) throughout core modules",
    "All exceptions caught and converted to AgentResponse — CLI never crashes",
    "Structured JSON logging per request (query, intent, tool, latency, source, error)",
    "Auto-clear __pycache__ on every startup ensures no stale bytecode",
    "Thread-safe singleton patterns (locks) in DraftManager, VectorStoreService, PermissionStore",
    "Configurable settings via environment variables / .env (pydantic-settings)",
    "LLM responses validated as strict JSON — free-text rejected to prevent misclassification",
]
for q in quality:
    add_bullet(q)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  11. BENEFITS & BUSINESS VALUE
# ════════════════════════════════════════════════════════════════════════════
add_heading("11. Benefits & Business Value", level=1)

add_heading("11.1  Productivity Benefits", level=2)
prod_benefits = [
    ("Instant document answers",       "Ask questions about PDFs, Word files, spreadsheets — get answers in seconds without opening any application."),
    ("Email triage automation",        "Semantic inbox search handles multi-field queries ('urgent emails from Alice about budget') that keyword search cannot."),
    ("AI reply drafting",              "Generate professional, grounded email replies in one command; review and send with explicit confirmation."),
    ("Voice meeting notes",            "Drop any audio recording; get full transcript + timestamped Q&A without manual note-taking."),
    ("Reminder management",            "Natural-language reminder creation with Windows notifications — no need to open calendar apps."),
    ("Context-aware follow-ups",       "Multi-turn conversation: 'summarise the above file', 'reply to it', 'what else is in that folder' all work without re-specifying."),
]
add_two_col_table(prod_benefits, header=["Capability", "Business Benefit"])

add_heading("11.2  Strategic Benefits", level=2)
strat = [
    ("100% data privacy",                "No documents, emails, or queries ever leave the machine. Compliant with internal data-handling policies."),
    ("Zero subscription cost",           "No cloud API fees. Runs on commodity hardware with a CPU-only Llama 3.2 1B model."),
    ("Extensible architecture",          "New agents can be added as Python modules. MCP server makes capabilities available to any AI client."),
    ("Works offline",                    "Fully functional without internet access — suitable for air-gapped or restricted environments."),
    ("Auditable AI actions",             "Every request is logged as a structured JSON line: intent, tool, latency, source, error."),
    ("No vendor lock-in",                "Models, embeddings, and vector databases are all open-source and swappable."),
]
add_two_col_table(strat, header=["Benefit", "Detail"])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  12. ARCHITECTURE MIND MAP  (text-based diagram)
# ════════════════════════════════════════════════════════════════════════════
add_heading("12. Architecture Mind Map", level=1)
add_body(
    "The diagram below shows the complete system as a hierarchical mind map, "
    "illustrating how every component relates to the central orchestrator."
)

diagram_lines = [
    "                       ┌──────────────────────────────┐",
    "                       │   LOCAL AI ASSISTANT v1.1    │",
    "                       │   (Central Orchestrator)      │",
    "                       └────────────┬─────────────────┘",
    "           ┌────────────────────────┼─────────────────────────────┐",
    "           │                        │                             │",
    "           ▼                        ▼                             ▼",
    "  ┌──────────────┐       ┌────────────────────┐       ┌─────────────────────┐",
    "  │  USER INPUT  │       │  INTENT PIPELINE   │       │  MEMORY & CONTEXT   │",
    "  ├──────────────┤       ├────────────────────┤       ├─────────────────────┤",
    "  │ CLI (main.py)│       │ 1. Regex fast-path │       │ last_file           │",
    "  │ Streamlit UI │       │ 2. Email context   │       │ last_folder         │",
    "  │ MCP Server   │       │ 3. Guardrails      │       │ last_intent         │",
    "  └──────────────┘       │ 4. LLM (Ollama)    │       │ last_email          │",
    "                         │ 5. Heuristic       │       │ user facts          │",
    "                         │ 6. Planner agent   │       │ conversation history│",
    "                         └────────┬───────────┘       └─────────────────────┘",
    "                                  │",
    "                     ┌────────────▼──────────────┐",
    "                     │  ACCESS CONTROL LAYER     │",
    "                     ├───────────────────────────┤",
    "                     │ Path extraction           │",
    "                     │ Permission store          │",
    "                     │ ALLOW_FOLDER / BLOCK      │",
    "                     │ Request permission        │",
    "                     └────────────┬──────────────┘",
    "                                  │",
    "          ┌──────────────────────▼──────────────────────┐",
    "          │                   ROUTER                     │",
    "          │  intent label  ──►  tool name                │",
    "          └────────────────────┬────────────────────────┘",
    "                               │",
    "    ┌──────────────────────────┼─────────────────────────────┐",
    "    │                          │                             │",
    "    ▼                          ▼                             ▼",
    "┌──────────────┐   ┌───────────────────────┐   ┌──────────────────────┐",
    "│  DOCUMENT    │   │  EMAIL SUBSYSTEM       │   │  UTILITY TOOLS       │",
    "│  SUBSYSTEM   │   ├───────────────────────┤   ├──────────────────────┤",
    "├──────────────┤   │ email.search           │   │ reminders.set/.list  │",
    "│ documents    │   │ → hybrid semantic+kw   │   │ → dateparser         │",
    "│ .search      │   │ email.reply            │   │ → plyer notifications│",
    "│ → RAG engine │   │ → LLM draft + anti-    │   │                      │",
    "│ → ChromaDB   │   │   hallucination prompt │   │ audio.transcribe     │",
    "│ → Sentence   │   │ email.send             │   │ → Whisper ASR        │",
    "│   Transformers│   │ → DraftManager         │   │ audio.query          │",
    "│ documents    │   │ → SMTP send_service     │   │ → ChromaDB audio     │",
    "│ .summarize   │   │ email.summarize         │   │                      │",
    "│ documents    │   │ → summarizer agent      │   │ system.chat          │",
    "│ .list        │   └───────────────────────┘   │ → Ollama LLM         │",
    "│ documents    │                               └──────────────────────┘",
    "│ .topics      │",
    "└──────────────┘",
    "",
    "  DATA STORES",
    "  ─────────────────────────────────────────────────────────────",
    "  data/vector_store_v2/     ← ChromaDB (documents)",
    "  data/vector_store_audio/  ← ChromaDB (audio transcripts)",
    "  data/vector_store_emails/ ← ChromaDB (email embeddings)",
    "  data/email_cache.json     ← Live IMAP email cache",
    "  data/drafts.json          ← Email draft persistence",
    "  data/reminders.json       ← Reminder persistence",
    "  data/memory.json          ← Conversation memory persistence",
    "  data/granted_folders.json ← Authorised folder list",
]

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("\n".join(diagram_lines))
run.font.name = "Courier New"
run.font.size = Pt(7.5)
run.font.color.rgb = C_NAVY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  13. FUTURE ROADMAP
# ════════════════════════════════════════════════════════════════════════════
add_heading("13. Future Roadmap", level=1)
add_body("Planned enhancements for upcoming versions:")

roadmap = [
    ("v1.2 — GPU acceleration",       "Enable CUDA inference in Ollama for sub-100ms LLM responses on GPU hardware."),
    ("v1.2 — LLM upgrades",           "Support larger models (llama3.2:8b, mistral) with automatic fallback to 1b on low-RAM machines."),
    ("v1.3 — Calendar integration",   "Connect to Outlook Calendar / Google Calendar for meeting scheduling via natural language."),
    ("v1.3 — Multi-user support",     "Separate memory and permission stores per user identity for shared machines."),
    ("v1.4 — Voice input",            "Microphone → Whisper STT → agent pipeline for fully voice-driven operation."),
    ("v1.4 — Web search agent",       "Optional online search agent (local proxy) for queries the LLM cannot answer from local data."),
    ("v2.0 — Agent-to-agent flow",    "Chained agents: e.g., 'summarise Q4 report and email it to Alice' as a single automated pipeline."),
    ("v2.0 — Fine-tuned model",       "Domain-specific fine-tuning of the base LLM on company documents for higher accuracy."),
]
add_two_col_table(roadmap, header=["Milestone", "Description"])

# ── Footer / signature ──────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─" * 60)
r.font.color.rgb = C_TEAL

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    f"Local AI Assistant v1.1  ·  Generated {datetime.date.today().strftime('%d %B %Y')}  ·  Confidential — Internal Use Only"
)
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ─────────────────────────────────────────────────────────────────────
output = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Local_AI_Assistant_Project_Report.docx")
doc.save(output)
print(f"✅  Document saved: {output}")
