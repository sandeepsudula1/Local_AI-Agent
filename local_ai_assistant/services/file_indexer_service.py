"""
services/file_indexer_service.py
=================================
SQLite-backed file metadata index for fast, path-independent file discovery.

Responsibilities
----------------
* Maintain a lightweight SQLite database of every file the system has ever
  been granted access to: path, name, size, mtime, extension, and a
  quick-scan content_hint (first 500 characters).
* Allow keyword search over file names AND content hints — no embeddings
  needed for name-based queries.
* Provide on-demand partial-text extraction (keyword + window) for FAST
  mode retrieval, avoiding full chunking when not needed.
* Expose a simple API consumed by the orchestrator, RAG engine, and file
  watcher service.

Database schema
---------------
  files(
      id          INTEGER PRIMARY KEY,
      path        TEXT UNIQUE,          -- absolute, normalised path
      name        TEXT,                 -- basename, lower-cased
      extension   TEXT,                 -- e.g. .txt
      size_bytes  INTEGER,
      mtime       REAL,                 -- os.path.getmtime epoch
      content_hint TEXT,               -- first 500 chars of raw text
      chunk_stale INTEGER DEFAULT 1,   -- 1 = needs re-chunk if used
      indexed_at  REAL                 -- epoch when row was inserted/updated
  )

Usage::

    from services.file_indexer_service import file_indexer

    # Register a folder (called automatically on permission grant)
    file_indexer.register_folder("/path/to/folder")

    # Search by name or content hint
    results = file_indexer.search("spring report")
    # → [{"path": "...", "name": "...", ...}, ...]

    # Get full metadata for one path
    meta = file_indexer.get_metadata("/path/to/spring_report.pdf")

    # Mark a file stale (called by file watcher on modification)
    file_indexer.mark_stale("/path/to/spring_report.pdf")

    # Remove a file (called by file watcher on deletion)
    file_indexer.remove("/path/to/spring_report.pdf")

    # Fast keyword-window extraction (FAST retrieval mode)
    snippet = file_indexer.keyword_window(
        "/path/to/report.txt", "deadline", window=400
    )
"""

from __future__ import annotations

import math
import os
import re
import sqlite3
import struct
import sys
import time
import threading
from pathlib import Path
from typing import Optional

from core.logging_config import get_logger
from configs.settings import settings, PROJECT_ROOT, DATA_DIR

log = get_logger(__name__)

_DB_PATH: Path = Path(DATA_DIR) / "file_index.db"

# Extensions recognised for content-hint extraction
_TEXT_EXTENSIONS: frozenset[str] = frozenset({
    ".txt", ".md", ".py", ".js", ".ts", ".java", ".json", ".csv", ".html",
    ".xml", ".yaml", ".yml", ".ini", ".cfg", ".log",
})

_CONTENT_HINT_LENGTH = 500   # characters to store as content_hint
_WINDOW_PADDING      = 200   # characters on each side of keyword match

# ── System scan exclusions ────────────────────────────────────────────────────
# Directories that must NEVER be scanned (system, hidden, or cache dirs)
_EXCLUDED_DIRS: frozenset[str] = frozenset({
    # Windows system
    "Windows", "Program Files", "Program Files (x86)", "ProgramData",
    "$Recycle.Bin", "System Volume Information", "Recovery",
    # User-level noise
    "AppData", ".git", ".hg", ".svn", "node_modules", "__pycache__",
    ".vscode", ".idea", "venv", "venv311", ".venv", "env",
    ".cache", ".npm", ".nuget", ".cargo", ".rustup",
    # Build artefacts
    "build", "dist", "target", "bin", "obj",
})

# Searchable file extensions for system scan
_SEARCHABLE_EXTENSIONS: frozenset[str] = frozenset({
    ".txt", ".md", ".py", ".js", ".ts", ".java", ".json", ".csv", ".html",
    ".xml", ".yaml", ".yml", ".ini", ".cfg", ".log",
    ".pdf", ".docx", ".doc", ".pptx", ".xlsx", ".xls",
    ".png", ".jpg", ".jpeg", ".webp", ".gif",
})

# Extensions that support richer text extraction (PDF, DOCX) beyond plain-text
_RICH_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx"})


class FileIndexerService:
    """SQLite-backed file metadata index."""

    def __init__(self, db_path: Path = _DB_PATH) -> None:
        self._db_path = db_path
        self._lock    = threading.Lock()
        self._init_db()

    # ── Schema ───────────────────────────────────────────────────────────────

    def _init_db(self) -> None:
        self._db_path.parent.mkdir(parents=True, exist_ok=True)
        with self._connect() as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS files (
                    id           INTEGER PRIMARY KEY AUTOINCREMENT,
                    path         TEXT    UNIQUE NOT NULL,
                    name         TEXT    NOT NULL,
                    extension    TEXT    NOT NULL DEFAULT '',
                    size_bytes   INTEGER NOT NULL DEFAULT 0,
                    mtime        REAL    NOT NULL DEFAULT 0.0,
                    content_hint TEXT    NOT NULL DEFAULT '',
                    chunk_stale  INTEGER NOT NULL DEFAULT 1,
                    indexed_at   REAL    NOT NULL DEFAULT 0.0
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS system_state (
                    key   TEXT PRIMARY KEY,
                    value TEXT
                )
            """)
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_name ON files(name)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_ext ON files(extension)"
            )
            # ── Migration: add content_embedding column if missing ───────────
            cursor = conn.execute("PRAGMA table_info(files)")
            col_names = {row["name"] for row in cursor.fetchall()}
            if "content_embedding" not in col_names:
                conn.execute(
                    "ALTER TABLE files ADD COLUMN content_embedding BLOB DEFAULT NULL"
                )
                log.info("[FileIndexer] Migrated: added content_embedding column")
        log.debug("[FileIndexer] DB ready at %s", self._db_path)

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(str(self._db_path), check_same_thread=False)
        conn.row_factory = sqlite3.Row
        return conn

    # ── Registration ─────────────────────────────────────────────────────────

    def register_file(self, path: str) -> bool:
        """Register or refresh a single file in the metadata index.

        Stores metadata, a quick content hint, and a content embedding
        (when text can be extracted).
        Returns True if the row was inserted/updated, False on error.
        """
        try:
            p = Path(path).resolve()
            if not p.is_file():
                return False
            stat      = p.stat()
            mtime     = stat.st_mtime
            size      = stat.st_size
            name      = p.name.lower()
            ext       = p.suffix.lower()
            hint      = self._extract_hint(p)
            now       = time.time()
            norm_path = str(p)

            # Build embedding from filename + content hint
            emb_blob = self._compute_embedding_blob(name, hint)

            with self._lock, self._connect() as conn:
                existing = conn.execute(
                    "SELECT mtime FROM files WHERE path = ?", (norm_path,)
                ).fetchone()
                if existing:
                    if abs(existing["mtime"] - mtime) < 0.5:
                        return True  # unchanged
                    conn.execute(
                        """UPDATE files
                           SET name=?, extension=?, size_bytes=?, mtime=?,
                               content_hint=?, content_embedding=?,
                               chunk_stale=1, indexed_at=?
                           WHERE path=?""",
                        (name, ext, size, mtime, hint, emb_blob, now, norm_path),
                    )
                    log.debug("[FileIndexer] Updated: %s", name)
                else:
                    conn.execute(
                        """INSERT INTO files
                           (path, name, extension, size_bytes, mtime, content_hint,
                            content_embedding, chunk_stale, indexed_at)
                           VALUES (?, ?, ?, ?, ?, ?, ?, 1, ?)""",
                        (norm_path, name, ext, size, mtime, hint, emb_blob, now),
                    )
                    log.debug("[FileIndexer] Registered: %s", name)
            return True
        except Exception as exc:
            log.warning("[FileIndexer] register_file(%r) failed: %s", path, exc)
            return False

    def register_folder(self, folder: str, recursive: bool = False) -> int:
        """Register all supported files in *folder*.

        Parameters
        ----------
        folder : str
            Absolute path to the folder.
        recursive : bool
            When True, scans subdirectories as well.

        Returns the number of files successfully registered.
        """
        count = 0
        try:
            root = Path(folder)
            if not root.is_dir():
                return 0
            iter_fn = root.rglob("*") if recursive else root.iterdir()
            for entry in iter_fn:
                if entry.is_file():
                    if self.register_file(str(entry)):
                        count += 1
        except Exception as exc:
            log.warning("[FileIndexer] register_folder(%r) failed: %s", folder, exc)
        log.info("[FileIndexer] Registered %d files from %r", count, folder)
        return count

    def scan_system(self, max_depth: int = 4, force: bool = False) -> int:
        """Scan user-accessible directories and build a searchable file index.

        Automatically detects search roots (user home, available drives),
        excludes system/hidden directories, and handles PermissionErrors
        gracefully.  Results are cached in the SQLite database for fast reuse.

        Parameters
        ----------
        max_depth : int
            Maximum directory depth to traverse from each root (default 4).
        force : bool
            If True, scan even if recently scanned.

        Returns the number of files successfully indexed.
        """
        if not force:
            with self._connect() as conn:
                res = conn.execute("SELECT value FROM system_state WHERE key='last_system_scan'").fetchone()
                if res:
                    try:
                        last_ts = float(res[0])
                        if time.time() - last_ts < 86400: # 24 hours
                            log.info("[FileIndexer] Skipping system scan (last scan < 24h ago)")
                            return 0
                    except (ValueError, TypeError):
                        pass

        # PART 2: CONTROLLED INDEXING — Logging
        print("[INDEX] Running in background...")
        log.info("[FileIndexer] Starting system scan (throttled)")

        # PART 1: DEFINE ROOT PATHS
        roots: list[str] = []
        user_home = os.path.expanduser("~")
        if os.path.isdir(user_home):
            roots.append(user_home)
            print(f"[FILE] Root: {user_home}")

        # Optionally include additional drives (D:\, E:\, etc.)
        if sys.platform == "win32":
            for letter in "DEFGHIJ":
                drive = f"{letter}:\\"
                if os.path.isdir(drive):
                    roots.append(drive)

        if not roots:
            log.warning("[FileIndexer] scan_system: no accessible roots found")
            return 0

        total = 0
        for root_path in roots:
            total += self._safe_walk_and_register(root_path, max_depth)

        # Update last scan time
        with self._connect() as conn:
            conn.execute("INSERT OR REPLACE INTO system_state (key, value) VALUES ('last_system_scan', ?)", (str(time.time()),))

        log.info("[FileIndexer] System scan complete: %d files indexed", total)
        print(f"[FILE] System scan complete: {total} files indexed")
        return total

    def _safe_walk_and_register(self, root: str, max_depth: int) -> int:
        """Walk a directory tree safely, skipping excluded and inaccessible dirs."""
        count = 0
        root_depth = root.rstrip(os.sep).count(os.sep)

        for dirpath, dirnames, filenames in os.walk(root):
            # PART 2: CONTROLLED INDEXING — Throttle to limit CPU usage
            time.sleep(0.02) # Yield to other processes

            # PART 4: PERFORMANCE — limit depth
            current_depth = dirpath.rstrip(os.sep).count(os.sep) - root_depth
            if current_depth >= max_depth:
                dirnames.clear()  # don't recurse deeper
                continue

            # PART 2: EXCLUDE SYSTEM FOLDERS — prune in-place
            dirnames[:] = [
                d for d in dirnames
                if d not in _EXCLUDED_DIRS
                and not d.startswith(".")
                and not d.startswith("$")
            ]

            # PART 3: SAFE DIRECTORY WALK — register files
            for fname in filenames:
                ext = os.path.splitext(fname)[1].lower()
                if ext not in _SEARCHABLE_EXTENSIONS:
                    continue
                fpath = os.path.join(dirpath, fname)
                try:
                    if self.register_file(fpath):
                        count += 1
                except PermissionError:
                    pass  # Skip inaccessible files silently
                except Exception:
                    pass  # Skip any other errors silently

        return count

    # ── Mutations ────────────────────────────────────────────────────────────

    def mark_stale(self, path: str) -> None:
        """Mark a file as stale so its chunks will be rebuilit on next use."""
        norm = str(Path(path).resolve())
        with self._lock, self._connect() as conn:
            conn.execute(
                "UPDATE files SET chunk_stale=1 WHERE path=?", (norm,)
            )
        log.debug("[FileIndexer] Marked stale: %s", norm)

    def remove(self, path: str) -> None:
        """Remove a file from the index (called when the file is deleted on disk)."""
        norm = str(Path(path).resolve())
        with self._lock, self._connect() as conn:
            conn.execute("DELETE FROM files WHERE path=?", (norm,))
        log.debug("[FileIndexer] Removed: %s", norm)

    def clear_stale_flag(self, path: str) -> None:
        """Clear the stale flag after successful re-chunk."""
        norm = str(Path(path).resolve())
        with self._lock, self._connect() as conn:
            conn.execute(
                "UPDATE files SET chunk_stale=0 WHERE path=?", (norm,)
            )

    # ── Search ───────────────────────────────────────────────────────────────

    def search(self, query: str, limit: int = 10) -> list[dict]:
        """Keyword search over file names and content hints.

        Returns a list of row dicts ordered by relevance score (descending).
        Score = 4 × (name matches) + 1 × (content_hint matches) for each
        query token.
        """
        tokens = [t.lower() for t in re.split(r"\s+", query.strip()) if len(t) >= 3]
        if not tokens:
            return []

        with self._connect() as conn:
            rows = conn.execute(
                "SELECT path, name, extension, size_bytes, mtime, "
                "content_hint, chunk_stale FROM files"
            ).fetchall()

        results: list[tuple[int, dict]] = []
        for row in rows:
            score = 0
            name_lc = row["name"]
            hint_lc = row["content_hint"].lower()
            for tok in tokens:
                if tok in name_lc:
                    score += 4
                if tok in hint_lc:
                    score += 1
            if score > 0:
                results.append((score, dict(row)))

        results.sort(key=lambda x: -x[0])
        
        # PART 1: RESULT CONFIDENCE
        # Max score assumes all tokens match in name (4 points each)
        max_possible = len(tokens) * 4 if tokens else 1
        
        # PART 3: FILE VALIDATION — Before returning file
        final_results = []
        for score, r in results:
            if len(final_results) >= limit:
                break
                
            fpath = r["path"]
            # check: exists and readable
            if os.path.exists(fpath) and os.access(fpath, os.R_OK):
                # Calculate confidence (0-100)
                conf = min(100, int((score / max_possible) * 100))
                r["confidence"] = conf
                final_results.append(r)
            else:
                # If not: auto-refresh index entry (remove stale entry)
                log.info("[FileIndexer] Auto-refresh: pruning invalid path: %s", fpath)
                self.remove(fpath)
        
        return final_results

    def semantic_search(self, query: str, limit: int = 10) -> list[dict]:
        """Embedding-based semantic search over indexed files.

        Embeds the query, computes cosine similarity against stored
        content_embedding BLOBs, and returns the top matches.

        Returns list of row dicts with an added ``semantic_score`` key.
        """
        try:
            from engines.embedding_engine import get_embedding_engine
            engine = get_embedding_engine()
            if not engine.is_ready:
                engine.load()
            if not engine.is_ready:
                return []

            q_vec = engine.embed(query, normalize=True)
            if not q_vec:
                return []
        except Exception as exc:
            log.debug("[FileIndexer] semantic_search – embed failed: %s", exc)
            return []

        with self._connect() as conn:
            rows = conn.execute(
                "SELECT path, name, extension, size_bytes, mtime, "
                "content_hint, chunk_stale, content_embedding FROM files "
                "WHERE content_embedding IS NOT NULL"
            ).fetchall()

        scored: list[tuple[float, dict]] = []
        dim = len(q_vec)
        for row in rows:
            vec = self._blob_to_vector(row["content_embedding"], dim)
            if vec is None:
                continue
            sim = self._cosine_similarity(q_vec, vec)
            d = dict(row)
            d.pop("content_embedding", None)
            d["semantic_score"] = round(sim, 4)
            scored.append((sim, d))

        scored.sort(key=lambda x: -x[0])
        log.debug(
            "[FileIndexer] semantic_search(%r) → top %d of %d rows",
            query, min(limit, len(scored)), len(scored),
        )
        return [d for _, d in scored[:limit]]

    def get_metadata(self, path: str) -> Optional[dict]:
        """Return stored metadata row for exact path, or None."""
        norm = str(Path(path).resolve())
        with self._connect() as conn:
            row = conn.execute(
                "SELECT * FROM files WHERE path=?", (norm,)
            ).fetchone()
        return dict(row) if row else None

    def lookup_by_name(self, name: str) -> list[dict]:
        """Find all registered files whose basename contains *name* (case-insensitive)."""
        name_lc = name.lower().strip()
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM files WHERE name LIKE ?",
                (f"%{name_lc}%",),
            ).fetchall()
        return [dict(r) for r in rows]

    def all_paths(self) -> list[str]:
        """Return all registered absolute paths."""
        with self._connect() as conn:
            rows = conn.execute("SELECT path FROM files").fetchall()
        return [r["path"] for r in rows]

    # ── Fast content extraction (FAST retrieval mode) ────────────────────────

    def keyword_window(
        self,
        path: str,
        keyword: str,
        window: int = _WINDOW_PADDING,
        max_windows: int = 3,
    ) -> str:
        """Extract up to *max_windows* text windows around *keyword* in the file.

        This is the FAST mode retrieval path — no chunking, no embeddings.
        Falls back to the content_hint stored in the DB when the file cannot
        be read directly.

        Parameters
        ----------
        path : str
            Absolute path to the file.
        keyword : str
            The term to locate in the file.
        window : int
            Characters on each side of each match to include.
        max_windows : int
            Maximum number of windows to return.

        Returns a single concatenated string of all found windows, or '' when
        the keyword is not found.
        """
        text = self._read_text(path)
        if not text:
            # Fall back to DB content hint
            meta = self.get_metadata(path)
            if meta:
                text = meta.get("content_hint", "")
        if not text:
            return ""

        kw_lc    = keyword.lower()
        text_lc  = text.lower()
        excerpts: list[str] = []
        start    = 0

        while len(excerpts) < max_windows:
            pos = text_lc.find(kw_lc, start)
            if pos == -1:
                break
            lo  = max(0, pos - window)
            hi  = min(len(text), pos + len(keyword) + window)
            excerpts.append(f"…{text[lo:hi]}…")
            start = hi

        return "\n\n".join(excerpts)

    def extract_partial(self, path: str, max_chars: int = 3000) -> str:
        """Read the first *max_chars* characters from a file (FAST mode fallback)."""
        text = self._read_text(path)
        return (text or "")[:max_chars]

    # ── Private helpers ──────────────────────────────────────────────────────

    def _compute_embedding_blob(self, name: str, hint: str) -> Optional[bytes]:
        """Return a BLOB of the normalised embedding for *name* + *hint*.

        Returns None when no text available or the engine is not loaded.
        Embedding engine is loaded lazily on first call.
        """
        text = f"{name} {hint}".strip()
        if len(text) < 3:
            return None
        try:
            from engines.embedding_engine import get_embedding_engine
            engine = get_embedding_engine()
            if not engine.is_ready:
                engine.load()
            if not engine.is_ready:
                return None
            vec = engine.embed(text, normalize=True)
            if not vec:
                return None
            return struct.pack(f"{len(vec)}f", *vec)
        except Exception as exc:
            log.debug("[FileIndexer] _compute_embedding_blob failed: %s", exc)
            return None

    @staticmethod
    def _blob_to_vector(blob: bytes, expected_dim: int) -> Optional[list[float]]:
        """Unpack a BLOB back into a list of floats."""
        if not blob:
            return None
        try:
            n = len(blob) // 4  # 4 bytes per float32
            if n != expected_dim:
                return None
            return list(struct.unpack(f"{n}f", blob))
        except Exception:
            return None

    @staticmethod
    def _cosine_similarity(a: list[float], b: list[float]) -> float:
        """Cosine similarity between two vectors (assumed pre-normalised)."""
        dot = sum(x * y for x, y in zip(a, b))
        # Vectors are normalised, but clamp for safety
        return max(0.0, min(1.0, dot))

    def _extract_hint(self, p: Path) -> str:
        """Read and return the first _CONTENT_HINT_LENGTH chars of a file.

        Supports plain text, PDF, and DOCX so that keyword + semantic search
        can match on content, not just filename.
        """
        ext = p.suffix.lower()
        try:
            if ext in _TEXT_EXTENSIONS:
                with p.open(encoding="utf-8", errors="ignore") as fh:
                    return fh.read(_CONTENT_HINT_LENGTH)
            if ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(str(p))
                text = "\n".join(
                    (page.extract_text() or "") for page in reader.pages
                )
                return text[:_CONTENT_HINT_LENGTH]
            if ext == ".docx":
                import docx
                doc = docx.Document(str(p))
                text = "\n".join(para.text for para in doc.paragraphs)
                return text[:_CONTENT_HINT_LENGTH]
        except Exception as exc:
            log.debug("[FileIndexer] _extract_hint(%r) failed: %s", str(p), exc)
        return ""

    @staticmethod
    def _read_text(path: str) -> str:
        """Best-effort full text extraction for common file types."""
        p = Path(path)
        if not p.is_file():
            return ""
        ext = p.suffix.lower()
        try:
            if ext in _TEXT_EXTENSIONS:
                return p.read_text(encoding="utf-8", errors="ignore")
            if ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(str(p))
                return "\n".join(
                    (page.extract_text() or "") for page in reader.pages
                )
            if ext == ".docx":
                import docx
                doc = docx.Document(str(p))
                return "\n".join(para.text for para in doc.paragraphs)
        except Exception as exc:
            log.debug("[FileIndexer] _read_text(%r) failed: %s", path, exc)
        return ""


# ---------------------------------------------------------------------------
# Global singleton
# ---------------------------------------------------------------------------

file_indexer = FileIndexerService()
